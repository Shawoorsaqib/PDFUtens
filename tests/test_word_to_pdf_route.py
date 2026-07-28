import io
import os
import unittest
import docx

from app import app
import config


class TestWordToPdfRoute(unittest.TestCase):

    def setUp(self):
        app.config["TESTING"] = True
        self.client = app.test_client()
        os.makedirs(config.UPLOAD_FOLDER, exist_ok=True)
        os.makedirs(config.OUTPUT_FOLDER, exist_ok=True)

    def _generate_dummy_docx_bytes(self):
        doc = docx.Document()
        doc.add_heading("Test Word Document", 0)
        doc.add_paragraph("This is a sample Word paragraph for PDF conversion testing.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def test_word_to_pdf_page_render(self):
        response = self.client.get("/word-to-pdf")
        self.assertEqual(response.status_code, 200)

    def test_word_to_pdf_upload_no_file(self):
        response = self.client.post(
            "/word-to-pdf/upload",
            data={},
            content_type="multipart/form-data"
        )
        self.assertEqual(response.status_code, 400)
        json_data = response.get_json()
        self.assertFalse(json_data["success"])
        self.assertIn("no file uploaded", json_data["message"].lower())

    def test_word_to_pdf_upload_invalid_extension(self):
        data = {
            "file": (io.BytesIO(b"Hello text content"), "sample.txt")
        }
        response = self.client.post(
            "/word-to-pdf/upload",
            data=data,
            content_type="multipart/form-data"
        )
        self.assertEqual(response.status_code, 400)
        json_data = response.get_json()
        self.assertFalse(json_data["success"])
        self.assertIn("unsupported file format", json_data["message"].lower())

    def test_word_to_pdf_upload_and_download_success(self):
        docx_buf = self._generate_dummy_docx_bytes()
        data = {
            "file": (docx_buf, "sample_document.docx")
        }
        response = self.client.post(
            "/word-to-pdf/upload",
            data=data,
            content_type="multipart/form-data"
        )
        self.assertEqual(response.status_code, 200)
        json_data = response.get_json()
        self.assertTrue(json_data["success"])
        self.assertIn("download_url", json_data)
        self.assertIn("filename", json_data)

        # Download the converted .pdf file
        download_url = json_data["download_url"]
        download_response = self.client.get(download_url)
        self.assertEqual(download_response.status_code, 200)
        self.assertEqual(
            download_response.mimetype,
            "application/pdf"
        )
        self.assertTrue(len(download_response.data) > 0)

        # Verify that output file was cleaned up from disk after downloading
        output_filename = json_data["filename"]
        output_path = os.path.join(config.OUTPUT_FOLDER, output_filename)
        self.assertFalse(os.path.exists(output_path))

    def test_download_nonexistent_file(self):
        response = self.client.get("/word-to-pdf/download/nonexistent_file.pdf")
        self.assertEqual(response.status_code, 404)


if __name__ == "__main__":
    unittest.main()
